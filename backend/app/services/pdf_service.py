import os
import subprocess
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document


def render_pdf(template_path: Path, data: dict, file_type: str) -> bytes:
    if file_type == "docx":
        return _render_docx(template_path, data)
    return _render_pdf_template(template_path, data)


def _replace_in_para(para, data: dict) -> None:
    """Replace placeholders in a paragraph, handling split runs."""
    full_text = "".join(run.text for run in para.runs)
    new_text = full_text
    for key, value in data.items():
        new_text = new_text.replace(f"{{{{{key}}}}}", str(value))
    if new_text != full_text and para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""


def _render_docx(template_path: Path, data: dict) -> bytes:
    doc = Document(template_path)

    for para in doc.paragraphs:
        _replace_in_para(para, data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, data)

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "doc.docx"
        pdf_path = Path(tmpdir) / "doc.pdf"
        doc.save(docx_path)

        env = {**os.environ, "HOME": "/tmp"}
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", tmpdir,
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            timeout=60,
            env=env,
        )

        return pdf_path.read_bytes()


def _fit_fontsize(text: str, font: fitz.Font, max_width: float, nominal_size: float, min_size: float = 5.0) -> float:
    """Scale font size down until text fits within max_width."""
    size = nominal_size
    while size > min_size:
        if font.text_length(text, fontsize=size) <= max_width:
            return size
        size -= 0.5
    return min_size


def _render_pdf_template(template_path: Path, data: dict) -> bytes:
    doc = fitz.open(str(template_path))

    for page in doc:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            areas = page.search_for(placeholder)
            if not areas:
                continue

            style = _find_placeholder_style(page, doc, placeholder)

            for area in areas:
                page.draw_rect(area, color=None, fill=(1, 1, 1))

                text = str(value)
                fontsize = _fit_fontsize(text, style["font"], area.width, style["size"])

                tw = fitz.TextWriter(page.rect)
                tw.append(
                    pos=(area.x0, area.y1 - fontsize * 0.15),
                    text=text,
                    font=style["font"],
                    fontsize=fontsize,
                )
                tw.write_text(page, color=style["color"])

    return doc.tobytes()


def _find_placeholder_style(page, doc, placeholder: str) -> dict:
    """Extract font, size, and color from the span that contains the placeholder."""
    blocks = page.get_text("rawdict")["blocks"]

    for block in blocks:
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                span_text = "".join(c.get("c", "") for c in span.get("chars", []))
                if placeholder not in span_text:
                    continue

                size = span.get("size", 11)
                color_int = span.get("color", 0)
                color = (
                    ((color_int >> 16) & 0xFF) / 255,
                    ((color_int >> 8) & 0xFF) / 255,
                    (color_int & 0xFF) / 255,
                )
                font = _load_pdf_font(doc, page, span.get("font", ""))
                return {"size": size, "color": color, "font": font}

    return {"size": 11, "color": (0, 0, 0), "font": _fallback_font("")}


def _load_pdf_font(doc, page, font_name: str) -> fitz.Font:
    """Try to extract the embedded font from the PDF; fall back to a system font."""
    clean = font_name.split("+")[-1].replace("-", "").replace(" ", "").lower()

    for font_info in page.get_fonts():
        xref = font_info[0]
        basefont = font_info[3]
        base = basefont.split("+")[-1].replace("-", "").replace(" ", "").lower()
        if clean and (clean == base or clean in base or base in clean):
            try:
                fd = doc.extract_font(xref)
                if fd and fd[3]:
                    return fitz.Font(fontbuffer=fd[3])
            except Exception:
                pass

    return _fallback_font(font_name)


def _fallback_font(font_name: str) -> fitz.Font:
    """DejaVu fonts support Cyrillic; choose serif/sans based on the original font name."""
    fn = font_name.lower()
    is_serif = any(s in fn for s in ["serif", "times", "roman", "palatino", "georgia", "garamond"])
    candidate = (
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf") if is_serif
        else Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
    )
    if candidate.exists():
        try:
            return fitz.Font(fontfile=str(candidate))
        except Exception:
            pass
    return fitz.Font("helv")
